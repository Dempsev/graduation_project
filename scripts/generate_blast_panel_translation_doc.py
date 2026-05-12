from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_DIR = Path(r"D:\graduation_project\coad\output\doc")
OUTPUT_PATH = OUTPUT_DIR / "毕业设计文献翻译-爆载荷拓扑优化.docx"
PAGE_DIR = Path(r"D:\graduation_project\coad\tmp\blast_pages")
CROP_DIR = Path(r"D:\graduation_project\coad\tmp\blast_crops")

TITLE = "拓扑优化在设计受爆载荷作用结构板中的应用"
SOURCE = (
    "Schiffer G, McMullen K, Bruhl J. Application of Topology Optimization to Design a Structural Panel "
    "Subjected to Blast Loading[C]//Proceedings of the ASME 2021 International Mechanical Engineering "
    "Congress and Exposition. 2021: IMECE2021-66667."
)

CONTENT = [
    {
        "level": 1,
        "title": "摘要",
        "paragraphs": [
            "本文旨在将拓扑优化方法应用于装甲战斗车辆底部防护板的设计，以应对简易爆炸装置（IED）造成的爆载荷威胁。过去二十年中，IED 的广泛使用导致大量士兵伤亡和车辆严重损毁，因此研究人员希望开发一种既能减小结构挠度、又能降低质量的防护板。",
            "研究目标是设计一种轻量化、模块化且成本可接受的结构板，以在提升乘员防护能力的同时改善车辆机动性。作者指出，随着增材制造技术的发展，拓扑优化可以通过减材式的材料分布策略生成高效而复杂的结构形式，这些结构往往难以通过传统制造手段实现。",
            "本文共进行了 8 组拓扑优化研究，得到多种不同的结构方案，并将优化结果与等质量、等高度的空心结构截面（HSS）等传统方案进行比较。比较指标包括最大挠度、应变能和应力。结果表明，优化结构的性能高度依赖于所采用的拓扑优化设计目标。文中提出的方法可为未来需要同时减小质量并提高刚度的设计任务提供参考。"
        ],
    },
    {
        "level": 1,
        "title": "1 引言",
        "paragraphs": [
            "自 21 世纪初以来，埋设爆炸物在伊拉克和阿富汗战场上造成了大量美军士兵死亡和受伤。这种严重后果推动了针对战斗车辆底部防护的持续研究，特别是车辆在遭受埋置弹药或简易爆炸装置攻击时的底部防护设计。理想的防护底板不仅应具有较高刚度、能够吸收爆炸能量以减小竖向挠度和车辆及乘员损伤，还应尽量减轻自身质量。",
            "相比传统厚重的一体式钢板或 V 形车底结构，轻量化、模块化和便携式防护板有望兼顾结构强度、机动性和维修便利性。模块化方案还可以缩短运输和抢修时间，提升车辆适用性。作者因此选择拓扑优化作为本研究的设计工具，希望开发出更轻、更硬的底部防护板结构。",
            "本文的核心目的并不是直接完成最终动态防爆结构设计，而是验证拓扑优化作为结构设计工具在此类问题中的有效性。虽然目标对象最终服役于动态爆载环境，但这项初步研究首先采用静载条件来比较不同拓扑优化设计的性能，以便量化拓扑优化相对于传统截面设计在刚度、强度和计算成本方面的潜在优势。"
        ],
    },
    {
        "level": 1,
        "title": "2 背景",
        "paragraphs": [
            "拓扑优化是一种材料分布方法，能够根据用户设定的设计目标，在给定设计空间内确定主要为各向同性材料的最佳布置方式。对于结构类问题，最常见的设计目标是以尽可能少的材料实现最小柔度，也就是在满足约束条件的前提下获得最大刚度。本文采用的拓扑优化以非线性有限元模型为基础，逐步判断各个单元对整体刚度的贡献，并保留那些对降低内部应变能最有效的材料区域。",
            "在优化开始之前，设计者需要明确载荷、边界条件以及质量或位移等约束。在本研究中，位移约束的作用是避免防护板挠度过大而伤及乘员，而质量约束则是为了保持车辆机动性。作者提到，优化过程中采用了 SIMP（带惩罚项的固体各向同性材料）方法，使模型更适合与各向同性材料及可制造性相匹配。",
            "作者还指出，拓扑优化与增材制造密切相关。由于优化得到的结构通常包含变化截面、小夹角和复杂几何，传统制造方式不易实现，而增材制造的逐层成形方式恰好适合这种复杂结构。即使无法直接采用增材制造，拓扑优化依然可以帮助设计者识别关键受力区域，再利用传统工艺制造与优化结果相近的结构。"
        ],
    },
    {
        "level": 2,
        "title": "2.1 拓扑优化设计目标",
        "paragraphs": [
            "本研究使用了两款商用拓扑优化软件：SolidWorks 和 LS-TaSC。两者都支持若干不同的设计目标，包括在质量约束下最小化位移、在位移约束下最小化质量，以及在质量约束下最大化刚度或刚重比。",
            "其中，最小位移设计目标会优先满足给定的减重比例，再在剩余材料中重新分配以减小挠度；最小质量设计目标则优先满足位移限制，只有在位移不超过约束值时才继续减小质量；最大刚度或刚重比设计目标则在施加给定减重比例的同时，尽量降低结构总应变能。作者后续正是通过比较这些不同目标下的结果，来分析拓扑优化方案性能变化的原因。"
        ],
    },
    {
        "level": 1,
        "title": "3 研究方法",
        "paragraphs": [
            "本文通过对不同几何截面开展拓扑优化研究，分析其性能变化规律。研究分别考察了截面几何参数和优化目标对优化结果的影响，并将拓扑优化得到的结构与传统实体截面、空心结构截面（HSS）和 I 形截面进行对比。所有对比均尽量保持其他变量不变，以便单独考察某一因素的影响。",
            "研究中共进行了 8 组对比试验。前三组用于探究减重比例、恒定面积和恒定惯性矩等几何约束对优化结构性能的影响，后五组则进一步考察不同软件和不同设计目标下优化结果的一致性，并重点与 HSS 和 I 形截面展开比较。",
            "几何约束、边界条件和载荷大小均参考了以往针对钛合金车辆底部防护板的小尺度研究。本文将梁截面视为整块防护板中的一个分析条带，全部有限元模型长度为 120 mm、宽度为 10 mm，两端均采用固支边界，以模拟底板通过螺栓固定在车体底部的情况。"
        ],
        "figures": [
            {"page": "page-03.png", "box": (620, 740, 1110, 1020), "caption": "图1 梁模型尺寸与边界支撑条件"}
        ],
    },
    {
        "level": 1,
        "title": "4 拓扑优化研究设计",
        "paragraphs": [
            "作者给出了 8 个拓扑优化研究案例，分别对应不同的设计目标、几何约束、载荷类型、边界条件和所用软件。材料统一采用 Ti-6Al-4V 钛合金，屈服强度为 827 MPa，弹性模量为 105 GPa，密度为 4430 kg/m^3。初始研究主要施加与 IED 压力水平相当的均布静载荷，以简化远距离爆炸作用下的底板受力问题，但此时未考虑爆载荷的动态效应。",
            "作者还使用第二款软件对静力分析结果进行交叉验证。虽然两款软件在材料、边界和载荷条件上保持一致，但网格类型和精度存在差别，第一款软件采用四面体网格，第二款软件则采用非线性实体四边形单元，并且由于计算条件限制，后者不得不使用较粗网格。这些差异也被纳入后续讨论。"
        ],
    },
    {
        "level": 1,
        "title": "5 结果",
        "paragraphs": [
            "研究结果按不同设计目标分别展开。作者首先分析了在“最大化刚重比”目标下，不同减重比例对优化梁性能的影响；随后又比较了在恒定面积和恒定惯性矩约束下，实体、HSS 和 I 形截面经优化后的性能变化；之后则重点针对 HSS 截面，比较了“最大化刚重比”“最小位移”“最小质量”等不同目标对结构表现的影响，并将 SolidWorks 与 LS-TaSC 的结果进行对照。"
        ],
    },
    {
        "level": 2,
        "title": "5.1 减重百分比：最大化刚重比",
        "paragraphs": [
            "第一组研究考察在“最大化刚重比”目标下，减重比例分别为 20%、40%、60% 和 80% 时优化梁的性能。结果显示，当减重比例为 20%、40% 和 60% 时，挠度相对增加幅度仍较有限；但当减重达到 80% 时，结构挠度显著放大，这是因为优化过程为了满足过高的减重要求，不得不移除一些对整体刚度非常关键的材料区域。"
        ],
        "figures": [
            {"page": "page-04.png", "box": (650, 480, 1110, 1080), "caption": "图2 不同减重比例下的优化梁结构"},
            {"page": "page-05.png", "box": (20, 20, 500, 360), "caption": "图3 质量与竖向挠度关系曲线"}
        ],
    },
    {
        "level": 2,
        "title": "5.2 恒定面积：最大化刚重比",
        "paragraphs": [
            "为了考察初始质量对“最大化刚重比”目标的影响，作者设计了恒定面积条件下的实体截面、HSS 截面和 I 形截面，并对其分别进行 60% 减重优化。结果表明，优化后实体截面出现了最小的相对挠度增幅，而 I 形截面挠度增幅最大。这说明，对于原本已具备较高强度效率的 I 形截面而言，继续从关键区域减材会显著削弱其抗弯刚度；而 HSS 的壁厚减薄同样会导致刚度下降。"
        ],
        "figures": [
            {"page": "page-05.png", "box": (20, 420, 630, 1020), "caption": "图4 恒定面积条件下的实体、HSS 和 I 形截面"},
            {"page": "page-05.png", "box": (700, 10, 1110, 420), "caption": "图5 恒定面积条件下各截面的质量与挠度对比"}
        ],
    },
    {
        "level": 2,
        "title": "5.3 恒定惯性矩：最大化刚重比",
        "paragraphs": [
            "为研究初始几何刚度对优化结果的影响，作者又构建了恒定惯性矩条件下的实体、HSS 和 I 形截面。为了获得相同的惯性矩，不同截面的初始质量并不相同。结果显示，在相同减重条件下，HSS 和 I 形截面的性能仍优于实体截面，但三者之间差异较前一研究更为复杂，这说明初始几何刚度对优化表现具有明显影响。"
        ],
        "figures": [
            {"page": "page-05.png", "box": (720, 430, 1115, 820), "caption": "图6 恒定惯性矩条件下的实体、HSS 和 I 形截面"},
            {"page": "page-06.png", "box": (20, 20, 560, 330), "caption": "图7 恒定惯性矩条件下各截面的质量与竖向挠度对比"}
        ],
    },
    {
        "level": 2,
        "title": "5.4 至 5.6 HSS 截面在不同设计目标下的比较",
        "paragraphs": [
            "作者进一步选取 HSS 截面作为代表，分别比较“最大化刚重比”“最小位移”和“最小质量”三种设计目标下的优化结果。在“最大化刚重比”目标下，优化后挠度仅略有下降，但应变能和最大应力都有上升，说明该目标并没有真正聚焦于位移最小化，而更偏向整体能量分布。",
            "在“最小位移”目标下，优化模型的整体变形模式发生改变，中跨挠度得到略微改善，但代价是应力显著集中于局部区域。作者据此认为，该目标在给定参考质量和减重比例条件下，主要优先保证质量约束，而不是纯粹追求最小挠度。",
            "在“最小质量”目标下，模型虽然试图在位移受限条件下继续减重，但结果显示位移增幅极大，已明显超出 HSS 的基准水平，说明该设计目标在本文所设约束组合下难以同时满足质量和位移要求。"
        ],
        "figures": [
            {"page": "page-06.png", "box": (20, 430, 560, 1030), "caption": "图8 10 mm 高 HSS 截面及其基于最大刚重比目标的优化结果"},
            {"page": "page-06.png", "box": (660, 390, 1110, 900), "caption": "图9 最小位移设计目标下 10 mm 高优化模型的变形形态"}
        ],
    },
    {
        "level": 2,
        "title": "5.7 商用拓扑优化软件对比",
        "paragraphs": [
            "为了比较两款商用软件在相同问题上的表现，作者在 LS-TaSC 中再次建立 HSS 与 I 形截面的优化模型，并采用位移载荷替代均布压力载荷，以适应软件的加载方式。结果显示，LS-TaSC 在最大刚度设计目标下能够更有效地降低 HSS 梁挠度，同时其应变能下降更明显、最大应力增幅更小；相比之下，SolidWorks 产生的结果则更容易出现局部应力集中。",
            "从机理上看，SolidWorks 的最大刚重比目标更接近于减少总应变能，但容易将材料移除到导致应力集中的区域；LS-TaSC 的多点设计方案则更强调应变能在结构内部的均匀分布，因此整体性能更平衡。作者同时指出，LS-TaSC 的计算成本显著更高，在该案例中求解时间约为 SolidWorks 的 120 倍。"
        ],
        "figures": [
            {"page": "page-07.png", "box": (600, 250, 1115, 900), "caption": "图10 10 mm 高截面在最大刚度设计目标下的等轴测图与立面图"},
            {"page": "page-08.png", "box": (620, 20, 1115, 560), "caption": "图11 两种拓扑优化软件在最大刚度目标下的结果对比"},
            {"page": "page-08.png", "box": (620, 560, 1115, 1030), "caption": "图12 10 mm 高优化模型的 von Mises 应力云图"}
        ],
    },
    {
        "level": 1,
        "title": "6 讨论",
        "paragraphs": [
            "作者认为，两款拓扑优化软件在设计目标的定义方式上存在本质差异。SolidWorks 所谓的“最大化刚重比”，更接近于在总体应变能意义上提高结构效率，但并不一定能保证局部应力或中跨挠度最优；LS-TaSC 则倾向于通过更均匀地分配材料来降低应变能密度，从而改善整体位移表现。",
            "此外，本文也指出了网格密度对优化结果的重要影响。随着网格加密，拓扑优化可能会移除更多微小单元，从而引入更细致的局部变形和应力集中。因此，未来在工程应用中需要结合网格敏感性分析，避免因网格过粗或过细而误判拓扑优化结构性能。"
        ],
    },
    {
        "level": 1,
        "title": "7 总结与未来工作",
        "paragraphs": [
            "本文在一系列静载梁模型上研究了拓扑优化在防爆结构板设计中的应用潜力。结果表明，拓扑优化确实能够为减重与增刚提供独特思路，但其具体效果高度依赖于设计目标设置。对于本文研究问题而言，更优的设计目标应能在减小质量的同时，直接聚焦于位移和应力控制，而不仅仅是降低总体柔度。",
            "作者进一步指出，要想真正将拓扑优化用于爆载荷防护板设计，还需要在未来引入动态爆载、考虑动态响应判据、分析非均匀压力分布，并研究不同设计目标与板件实际防爆性能之间的关系。例如，最小化应力集中是否能防止板件破坏、最小化竖向挠度是否真正降低乘员伤害，都需要进一步验证。",
            "总体而言，本文证明了拓扑优化可以作为设计受爆载荷结构板的前期探索工具，为后续动态防爆结构研究奠定了方法基础。"
        ],
    },
]


def set_run_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_blank_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)


def add_heading(doc: Document, title: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 16
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 15 if level == 2 else 14
    run = p.add_run(title)
    set_run_font(run, "黑体", size, True)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, "宋体", 12, False)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "宋体", 10, False)


def crop_image(page_name: str, box: tuple[int, int, int, int]) -> Path:
    CROP_DIR.mkdir(parents=True, exist_ok=True)
    page_path = PAGE_DIR / page_name
    crop_path = CROP_DIR / f"{Path(page_name).stem}-{'-'.join(map(str, box))}.png"
    if not crop_path.exists():
        with Image.open(page_path) as img:
            img.crop(box).save(crop_path, format="PNG")
    return crop_path


def add_figure(doc: Document, figure: dict) -> None:
    image_path = crop_image(figure["page"], figure["box"])
    with Image.open(image_path) as img:
        w, h = img.size
    width_cm = min(14.2, w / 95)
    if h > w * 1.2:
        width_cm = min(10.6, width_cm)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, figure["caption"])


def build_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(TITLE)
    set_run_font(run, "黑体", 18, True)

    add_blank_line(doc)

    for block in CONTENT:
        add_heading(doc, block["title"], block["level"])
        for para in block["paragraphs"]:
            add_body_paragraph(doc, para)
        for fig in block.get("figures", []):
            add_figure(doc, fig)
        add_blank_line(doc)

    for _ in range(4):
        add_blank_line(doc)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run("译文原文出处：")
    set_run_font(run, "黑体", 12, True)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(SOURCE)
    set_run_font(run, "宋体", 12, False)
    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
