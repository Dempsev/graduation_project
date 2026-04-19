from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


OUTPUT_DIR = Path(r"D:\graduation_project\coad\output\doc")
OUTPUT_PATH = OUTPUT_DIR / "毕业设计文献翻译.docx"
IMAGE_DIR = Path(r"D:\graduation_project\coad\tmp\pdf_images")
NORMALIZED_IMAGE_DIR = IMAGE_DIR / "normalized"


TITLE = "基于深度学习的高自由度声学超材料定制化吸声逆向设计"
SOURCE = (
    "Yan, J., Li, Y., Yin, G., Yao, S., Peng, Y. "
    "Inverse design on customised absorption of acoustic metamaterials with high degrees "
    "of freedom by deep learning[J]. Mechanical Systems and Signal Processing, 2025, 237: 112989. "
    "DOI: 10.1016/j.ymssp.2025.112989."
)


CONTENT = [
    {
        "level": 1,
        "title": "摘要",
        "paragraphs": [
            "由于传统试错式设计需要反复迭代并消耗大量计算与实验资源，声学超材料按需设计一直较为困难。本文提出了一种基于条件生成对抗网络（CGAN）的逆向设计方法，用于带穿孔板的多孔超材料（MMPP）定制化吸声结构设计。该方法能够在高自由度设计空间中，快速生成嵌入多孔基体内部的刚性夹杂构型。",
            "在数据集构建阶段，作者结合已有声学设计知识，构造了包含封闭刚性夹杂和薄壁夹杂两类样本的自定义数据集，使其同时具备宽频和低频吸声潜力。训练完成后，CGAN 学习了目标宽频吸声谱与 MMPP 构型之间的非线性映射关系。测试结果表明，逆向设计得到的预测吸声曲线与目标曲线之间具有较高一致性，平均均方误差为 5.03×10^-3，平均绝对百分比误差为 6.39%。",
            "与传统纯数据驱动方法不同，本文提出的方法将结构约束与基于物理规律的后处理步骤相结合，因此能够生成可制造、可解释并且超出训练样本分布的新型结构。不同声学性能对应的生成图样反映出较清晰的物理关联：面向低频吸声优化的结构通常具有更低的填充率和更浅层的夹杂嵌入位置，而面向高频吸声优化的结构则往往具有更高的填充率和更深层的夹杂分布。",
            "此外，该方法可以针对用户自定义的吸声目标输出多组候选方案。在总体厚度为 51 mm 的条件下，逆向设计所得结构可在 360-3000 Hz 范围内实现有效吸声（α>0.5）。其多重吸收峰来源于微穿孔共振、嵌入夹杂的局部共振以及高阶模态共振与声波干涉的共同作用。该研究为声学超材料的快速性能驱动设计提供了新的技术路线，也为揭示结构-性能耦合机理提供了新的思路。"
        ],
    },
    {
        "level": 1,
        "title": "1 引言",
        "paragraphs": [
            "声学超材料因其独特的声学特性，近年来成为噪声控制领域的重要研究方向。传统多孔吸声材料虽然能够较好地耗散中高频声能，但在紧凑结构中实现低中频高效吸声依然较为困难。为改善这一问题，研究者提出了多孔超材料的概念，即在多孔基体中嵌入刚性夹杂，使材料黏滞损耗、热损耗与附加的局部共振或散射机制共同作用，从而提升整体吸声性能。已有研究表明，夹杂的形状、尺寸及空间位置对吸声特性具有显著影响，因此相关结构必须进行精细设计。",
            "传统设计方法主要依赖研究者经验以及有限元法、传递矩阵法等正向分析工具，通过不断调整参数并重复仿真来逼近目标性能。虽然这种方法已经取得了不少成功应用，但它天然存在过程繁琐、周期较长以及对经验依赖强等不足。同时，设计变量数量与设计空间范围也受到限制。由于不同子结构对吸声性能的影响往往是耦合的，逐个参数独立优化也未必能够得到全局最优解。",
            "为了提高设计效率，研究者将遗传算法、粒子群优化等优化方法引入声学超材料设计中，能够在既定变量范围内搜索较优结构参数。这些方法在多层谐振器、镍泡沫复合吸声结构等研究中都表现出较好的效果。然而，当超材料几何形态复杂、材料分布精细且参数维度很高时，搜索空间将急剧扩大，优化算法的表现容易受到搜索空间特征和算法超参数的影响，进一步提升设计效率仍然十分必要。",
            "随着计算硬件和深度学习技术的发展，越来越多研究开始尝试将数据驱动方法用于超材料逆向设计。逆向设计的核心思想是：直接从目标动态性能出发，反推出满足要求的结构构型，而不必沿用传统的物理启发式正向设计流程。与传统方法相比，逆向设计通常可以探索更广阔的设计空间，但其前提是需要建立合适的结构表达方式，并构建足够大的训练数据集。",
            "当前结构表征方式大致有两类。一类是利用若干几何参数定义结构，再通过系统改变这些参数建立样本库；另一类是将结构像素化并用矩阵表示材料分布。这种像素化方式更适合表达高自由度复杂结构，但同时也提高了模型学习难度。因此，对于高维设计任务，需要采用更适合处理高维图像与分布映射问题的网络模型，例如生成对抗网络（GAN）及其条件扩展形式 CGAN。",
            "本文面向吸声性能按需设计问题，提出了基于 CGAN 的 MMPP 逆向设计方法。作者首先借助已有声学知识构造兼顾低频和宽频吸声能力的数据集，然后通过图像连通域分析、滤波和平滑等手段对输入输出图样进行处理。全文随后从数据集构建、网络架构、模型训练、结果分析、实验验证和结论六个方面展开。"
        ],
    },
    {
        "level": 1,
        "title": "2 数据集构建",
        "paragraphs": [
            "按需设计声学超材料的关键在于构造能够覆盖目标吸声频段的数据集。为此，本文引入三条先验设计原则：其一，选用低或中等流阻率的多孔材料作为宽频吸声基元；其二，在多孔材料中嵌入刚性夹杂，形成多孔超材料，以提升整体吸声能力；其三，在结构表面引入穿孔板，以增强低频吸声效果。为了兼顾设计自由度与计算可行性，作者采用数字材料方法，将连续材料域离散为有限的二值单元。",
        ],
    },
    {
        "level": 2,
        "title": "2.1 数字图样生成策略",
        "paragraphs": [
            "本文所设计的 MMPP 主要由两部分组成：一是嵌入刚性夹杂的多孔超材料，二是覆盖于其上的穿孔板。数据集样本采用 88×88 像素的二值图像表示结构截面，其中“0”表示多孔材料，“1”表示刚性夹杂。该分辨率既能表达较复杂的结构特征，又能避免过高分辨率带来的计算成本激增。对于尺寸固定的多孔基体而言，夹杂的形状与尺寸是决定吸声性能的关键因素，因此数据集必须尽可能覆盖多样且具有代表性的图样。",
            "作者在数据集中设置了两类夹杂：封闭刚性夹杂与薄壁刚性夹杂。所有样本的生成都满足可制造性与功能性约束，避免出现难以加工或失去声学意义的异常结构。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-003.jpg",
                "caption": "图1 0/1 分布矩阵及对应数字图样；不同 c1 与 c2 条件下嵌入封闭刚性夹杂的多孔材料示意图"
            }
        ],
    },
    {
        "level": 3,
        "title": "2.1.1 多孔材料中的封闭刚性夹杂",
        "paragraphs": [
            "封闭刚性夹杂是指没有开口的封闭形状。已有研究表明，这类夹杂在多孔超材料中的填充率和质心位置会显著影响吸声频率。为灵活控制夹杂与多孔材料之间的边界，作者采用闭合 B 样条曲线描述夹杂轮廓，并在极坐标下用一组参数控制其位置、尺寸和形状。其中，参数 h 用于调节夹杂在多孔基体中的质心位置，r0 用于控制包围区域整体尺度，c1 与 c2 则共同决定夹杂的具体外形。",
            "为了获得丰富的样本，作者将 h、r0、c1 和 c2 在合理范围内离散化并进行系统组合，从而快速批量生成不同填充率、不同位置和不同形状的夹杂图样。这种参数化图样生成方式既能保证设计空间的覆盖度，也便于后续深度学习模型训练。"
        ],
    },
    {
        "level": 3,
        "title": "2.1.2 多孔材料中的薄壁刚性夹杂",
        "paragraphs": [
            "对于薄壁刚性夹杂，作者在物理约束范围内利用随机函数生成控制参数，再将这些参数组合成图样。薄壁夹杂包含三类基本元件：隔板、矩形腔体和 C 形腔体。为了增加图样多样性，这些元件既可以单独出现，也可以两两组合，从而得到多种具有局部共振特性的结构形式。",
            "由于此类夹杂的吸声能力与局部共振密切相关，隔板间距、腔体开口尺寸和整体腔体尺度都必须受到约束。作者设置了边界留白、壁厚、隔板最小间距以及腔体尺度与开口尺寸等约束条件，以保证生成图样在加工、打印和实际使用中的可行性。总体而言，这一类夹杂主要通过改变局部共振模式来调节吸声频率。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-004.jpg",
                "caption": "图2 含隔板、矩形腔体和 C 形腔体等薄壁刚性夹杂的典型多孔材料图样"
            },
            {
                "path": IMAGE_DIR / "img-005.jpg",
                "caption": "图3 图样预处理示意图：孤立元件与连通元件两种情况"
            }
        ],
    },
    {
        "level": 2,
        "title": "2.2 数字图样吸声谱的获取",
        "paragraphs": [
            "仅有结构图样还不足以支撑逆向设计模型训练，还必须为每一张图样匹配对应的吸声谱。因此，作者采用理论分析与有限元数值计算相结合的方式，建立每个样本在 40-3000 Hz 频段内的吸声系数数据。"
        ],
    },
    {
        "level": 3,
        "title": "2.2.1 吸声理论基础",
        "paragraphs": [
            "MMPP 的理论模型由穿孔板和多孔超材料两部分组成。穿孔板通过其比声阻抗体现对声场的影响，多孔材料部分则采用 Johnson-Champoux-Allard（JCA）模型进行表征。该模型利用孔隙率、迂曲度、流阻率、黏性特征长度和热特征长度等关键参数，描述多孔材料的有效动态质量密度与体积模量，进而推导整体结构的声学阻抗、反射特性及吸声系数。",
            "这一理论框架说明了 MMPP 的吸声来源并非单一机制，而是穿孔板共振、多孔介质黏滞与热耗散以及夹杂引发的局部共振共同作用的结果。理论分析也为后续有限元建模与结果解释提供了基础。"
        ],
    },
    {
        "level": 3,
        "title": "2.2.2 基于数值模型的吸声谱计算",
        "paragraphs": [
            "在获得图样之后，作者利用 COMSOL Multiphysics 6.0 中的有限元方法对不同结构的吸声谱进行批量计算。由于样本最初以像素图形式存在，因此需先将图像转换为可被有限元软件识别的几何模型。文中使用 MATLAB 进行图像预处理和命令自动化，并借助 COMSOL 与 MATLAB 的 LiveLink 模块，实现从图样到仿真的自动流程。",
            "考虑到结构沿竖向截面保持一致，作者采用二维有限元模型替代三维模型，以显著降低计算成本。通过与理论结果比较，二维模型在关键吸声特征上具有较高可信度，因此能够用于大规模样本仿真与数据集生成。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-006.jpg",
                "caption": "图4 二维有限元模型及其网格划分示意图"
            },
            {
                "path": IMAGE_DIR / "img-007.jpg",
                "caption": "图5 数值模型验证及数据集中典型带穿孔板结构的吸声谱与图样示例"
            }
        ],
    },
    {
        "level": 2,
        "title": "2.3 数据集组成",
        "paragraphs": [
            "最终，作者共收集了 14420 个 88×88 像素的数字图样，并分别计算了每个图样在 75 个频点（40-3000 Hz）上的吸声系数。整个数据集由结构图像和对应的吸声谱共同组成，且在不同夹杂类型之间尽量保持均衡，以提高模型训练的稳定性与泛化能力。",
            "数据分布分析表明，这些样本在 256-3000 Hz 范围内普遍可以实现有效吸声（α>0.5），覆盖了交通噪声、施工噪声和一般工程噪声中的主要频率成分。虽然数据集构建需要大量有限元仿真，但这是一种一次性计算投入；一旦后续模型训练完成，结构预测就可以在数秒内完成，因此具有明显的工程价值。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-008.jpg",
                "caption": "图6 数据集获取流程以及二维 MMPP 样本的吸声系数分布"
            }
        ],
    },
    {
        "level": 1,
        "title": "3 用于 MMPP 逆向设计的 CGAN",
        "paragraphs": [
            "本文逆向设计方法的核心是构建并训练 CGAN，使其同时学习结构图样的分布规律以及图样与吸声谱之间的非线性映射关系。"
        ],
    },
    {
        "level": 2,
        "title": "3.1 CGAN 的工作流程与机理",
        "paragraphs": [
            "传统 GAN 由生成器 G 和判别器 D 构成，通过两者之间的对抗训练来学习真实样本分布。经典 GAN 主要用于无条件图像生成，而 CGAN 则在此基础上引入额外条件，从而能够根据给定信息生成对应结果。对于本文而言，这一额外条件就是目标吸声谱，因此 CGAN 特别适合用于按目标声学性能反推结构构型的任务。",
            "训练过程中，首先初始化 G 与 D 的参数；随后从训练集中提取一批样本。生成器依据吸声谱条件和随机噪声生成候选图像，判别器则同时接收真实图像及其对应吸声谱，以及生成图像及其条件信息，并判断输入属于“真实”还是“伪造”。在固定 G 的条件下不断更新 D，当 D 具有较强区分能力后，再固定 D、通过反向传播更新 G。随着训练持续进行，G 的“欺骗能力”和 D 的“辨别能力”逐渐达到平衡，最终形成稳定的条件生成能力。文中采用二元交叉熵损失函数作为对抗训练目标。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-009.jpg",
                "caption": "图7 GAN 与 CGAN 的训练流程示意图"
            }
        ],
    },
    {
        "level": 2,
        "title": "3.2 CGAN 的网络结构",
        "paragraphs": [
            "生成器和判别器都基于卷积神经网络构建，因为卷积结构尤其擅长处理二维图像数据。生成器包含两个输入通道：一条接收 1×75 维的吸声谱，另一条接收 1×128 维高斯噪声。噪声的引入使得即使面对相同目标谱，生成器也能输出多样化的候选结构。两个输入在分别经过全连接层和重塑之后被融合，并依次通过五层卷积层和上采样层，最终输出 88×88 像素的二值化结构图像。",
            "在激活函数方面，生成器前四层采用 LeakyReLU，最后一层采用 Tanh，以缓解梯度消失问题并使输出范围与预处理后的图像数据保持一致。判别器同样具有两个输入通道，一条接收待分类图像，另一条接收对应吸声谱并将其重塑为二维形式后与图像拼接。拼接后的数据经过四层卷积层处理后，最终通过 Sigmoid 输出“真实”或“伪造”的概率。作者强调，若 G 与 D 的深度不平衡，就可能出现模式崩溃或判别器过拟合，因此网络架构是在综合计算资源与已有研究经验后反复试验确定的。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-010.jpg",
                "caption": "图8 CGAN 详细网络结构：生成器与判别器示意图"
            }
        ],
    },
    {
        "level": 2,
        "title": "3.3 CGAN 的训练细节",
        "paragraphs": [
            "为了提高训练效率和模型稳定性，作者对图像和吸声谱数据都进行了专门预处理，并仔细设置了优化器、学习率、批大小等超参数。"
        ],
    },
    {
        "level": 3,
        "title": "3.3.1 数据预处理",
        "paragraphs": [
            "原始二值图像的取值范围为 [0,1]。考虑到大量 0 值不利于网络首层充分激活，作者将像素值变换为 [-1,1]，即将原来的 0 替换为 -1，以加快收敛速度。同时，这一处理也与生成器末层 Tanh 的对称输出范围保持一致。",
            "此外，作者对 14420 个样本进行随机打乱，以避免样本顺序带来的偏置，然后按 0.85:0.15 的比例划分训练集和测试集。吸声谱数据还进行了标准化处理，以提高训练效率并降低过拟合风险。"
        ],
    },
    {
        "level": 3,
        "title": "3.3.2 超参数设置",
        "paragraphs": [
            "模型训练采用 Adam 优化器，其中生成器学习率设为 0.0001，判别器学习率设为 0.0004，用以平衡两者的收敛速度。训练总迭代次数为 10000 次，批大小为 32。生成器五层卷积层的滤波器数量依次为 128、128、64、32 和 1；判别器四层卷积层的滤波器数量依次为 8、16、32 和 64。为防止判别器过拟合，作者在其各层加入了 60% 的 dropout 正则化。",
            "训练损失曲线显示，在前 1000 次迭代中模型波动较明显，随后逐步趋于稳定。最终，生成器损失大致稳定在 3-5 之间，判别器对真实和伪造样本的损失都逼近 0，说明对抗训练达到相对均衡状态并取得预期效果。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-011.jpg",
                "caption": "图9 图像数据预处理示意图及 CGAN 训练损失曲线"
            }
        ],
    },
    {
        "level": 1,
        "title": "4 结果与讨论",
        "paragraphs": [
            "在 CGAN 训练完成后，模型既能识别输入图样的统计特征，又能学习吸声谱与结构之间的映射关系。将目标吸声曲线输入训练好的模型后，即可生成相应的 MMPP 候选结构。"
        ],
    },
    {
        "level": 2,
        "title": "4.1 图像后处理",
        "paragraphs": [
            "由于 88×88 像素图样具有很高的自由度，模型直接生成的结构中经常包含孤立小像素块、极细连接和锯齿边界。这些细节不利于实际加工，也会增加制造成本。因此，作者对生成图像执行两类后处理：一是去除冗余的小像素点和不必要的微小连通域，二是对材料边界进行平滑处理。",
            "为验证后处理不会破坏声学性能，作者比较了三个随机样本在处理前后的吸声谱，结果显示二者几乎重合，说明这些几何微扰对总体吸声性能影响极小。也就是说，后处理既提升了可制造性，又保持了原有设计性能。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-012.jpg",
                "caption": "图10 CGAN 生成图样的后处理过程及处理前后吸声性能对比"
            }
        ],
    },
    {
        "level": 2,
        "title": "4.2 训练后 CGAN 在测试集上的评估",
        "paragraphs": [
            "为了定量评估逆向设计精度，作者将测试集中从未参与训练的 2163 条吸声谱输入训练好的 CGAN，并对生成结构进行后处理和有限元复算，再将预测得到的吸声谱与目标谱进行比较。评估指标采用均方误差（MSE）和平均绝对百分比误差（MAPE）两项，以避免仅依赖单一绝对误差或相对误差带来的偏差。",
            "结果表明，在整个频率范围内，平均 MSE 约为 5.03×10^-3，平均 MAPE 约为 6.39%。虽然低频段因目标吸声系数本身较小而导致相对误差偏高，但其绝对误差仍处于较低水平，因此总体预测精度是令人满意的。",
            "除了精度之外，作者还考察了生成图样的有效性与多样性。由于原始数据集中并不包含完全均匀、没有夹杂的多孔材料，因此若后处理后退化为无夹杂结构，则视为无效设计。2163 个测试样本中仅有 2 个样本属于这种情况，图样有效率达到 99.91%。同时，模型生成的结构可归纳为 12 种典型类型，说明 CGAN 不仅能复现已有样式，也能在训练分布附近产生多样化设计。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-013.jpg",
                "caption": "图11 测试集中目标吸声谱与生成结构预测吸声谱之间的 MSE、MAPE 及典型图样类型"
            },
            {
                "path": IMAGE_DIR / "img-014.jpg",
                "caption": "图12 训练后模型设计图样的预测吸声谱与测试集目标吸声谱对比"
            }
        ],
    },
    {
        "level": 2,
        "title": "4.3 MMPP 的逆向设计性能",
        "paragraphs": [
            "为了更直观地展示模型性能，作者从测试集中随机选取了 9 条吸声谱作为目标曲线，并比较训练后模型生成结构的预测吸声谱与目标曲线之间的差异。除采用 MSE 衡量点对点误差外，文中还引入余弦相似度 k 来反映两条曲线在整体形状上的一致性，k 越接近 1，说明曲线形状越相似。",
            "结果显示，这 9 组样本的余弦相似度均非常接近 1，MSE 也普遍接近 0。其中某一组样本的余弦相似度达到 0.9999，MSE 仅为 1.288×10^-4，几乎与目标曲线完全重合。除个别样本在局部峰值处存在稍大偏差外，整体趋势都与目标高度一致，表明训练后的 CGAN 具有较强的逆向设计能力。"
        ],
    },
    {
        "level": 2,
        "title": "4.4 CGAN 对物理特征的隐式学习",
        "paragraphs": [
            "为了进一步提升逆向设计的可解释性，作者选取在约 700 Hz 和 2400 Hz 处具有双峰的目标吸声曲线，并将其划分为三类：第一类强调低频峰值、第二类强调双峰都较强的宽频设计、第三类强调高频峰值。针对每一类目标曲线，CGAN 都生成了三组候选结构，并且预测吸声曲线均与目标良好匹配，说明模型在保持精度的同时具备较强的设计多样性。",
            "更重要的是，不同目标曲线对应的结构特征显示出清晰的物理规律。低频吸声优化结构通常采用尺寸较小、填充率较低且埋设较浅的夹杂；宽频优化结构则具有中等填充率、形态较多样并集中分布在中部区域；高频吸声优化结构往往使用体积更大、组合更复杂且埋设更深的夹杂，填充率也显著升高。",
            "这一结果说明，CGAN 并不是单纯进行“图像拟合”，而是在数据驱动的训练过程中隐式学习了吸声性能与结构配置之间的内在联系，从而能够为设计者提供更具物理意义的结构启发。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-015.jpg",
                "caption": "图13 CGAN 生成的结构构型及其对应的预测吸声曲线"
            }
        ],
    },
    {
        "level": 2,
        "title": "4.5 MMPP 的定制化低频宽带吸声",
        "paragraphs": [
            "作者将基于 CGAN 的定制化逆向设计流程总结为五个步骤：首先根据应用场景定义目标吸声谱；然后将目标谱输入训练好的 CGAN；接着对生成结构进行后处理，以保证其具备实际制造可行性；之后通过有限元和实验对结构性能进行验证；最后在多组候选方案中选择误差最小或更便于制造的设计付诸实施。",
            "为了展示该流程的实际能力，作者定义了一条低频宽带吸声目标曲线：在 360 Hz 处吸声系数达到 0.5，在 560 Hz 以上达到 0.8。基于这一目标，训练好的 CGAN 生成了四组候选设计，其预测吸声谱均能较好满足给定目标。这说明即使面对同一条目标曲线，模型也可以给出多种结构方案，为工程实现提供更大的选择余地。",
            "与依赖复杂物理方程反复求解的拓扑优化不同，CGAN 只需要进行一次离线训练，完成训练后即可在秒级时间内输出候选结构。虽然前期数据集构建和模型训练仍需要较大计算投入，但在后续大量重复设计任务中，数据驱动方法在效率、成本和可复用性方面都具有明显优势。作者同时指出，模型生成的所有候选结构仍需经过物理验证，领域知识仍然是最终设计决策的重要依据。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-016.jpg",
                "caption": "图14 基于 CGAN 的 MMPP 按需逆向设计流程示意图"
            },
            {
                "path": IMAGE_DIR / "img-017.jpg",
                "caption": "图15 四种候选设计的目标吸声谱、预测吸声谱及对应逆向设计结构"
            }
        ],
    },
    {
        "level": 2,
        "title": "4.6 MMPP 的吸声机理",
        "paragraphs": [
            "为了揭示四组逆向设计方案的吸声机制，作者选取误差最小的 Design-1 作为代表，并对二维有限元模型与三维有限元模型的结果进行比较。结果显示，两种模型得到的吸声谱基本一致，说明二维模型能够较好反映结构的主要吸声特征。随后，作者利用三维模型进一步分析非共振频率和三个吸收峰频率处的声压场与声速场分布。",
            "在第一个吸收峰 P1 处，声压主要分布于穿孔板后方的多孔材料区域，说明此时的吸声主要来自穿孔板与后部多孔材料形成的复合腔体结构。声波通过穿孔板后与后方吸声层发生有效耦合，穿孔口附近空气质量产生类质量-弹簧共振，并在多孔材料的黏滞和热耗散作用下实现低频吸收增强。",
            "在第二个吸收峰 P2 处，薄壁夹杂（如 C 形腔、矩形腔和隔板等）形成的局部共振系统起主导作用。声压和声速矢量集中在夹杂附近，表明该频率下局部能量高度集中，邻近多孔材料中的耗散能力随之增强，从而实现较强吸收。",
            "在第三个吸收峰 P3 处，声压主要集中在多孔区域上部及夹杂边界附近，并伴随更密集的声速扰动与回传现象。由于高频声波波长更短，夹杂边界处更容易发生多重散射与干涉；在高阶模态耦合作用下，多孔材料内的黏滞和热损耗被更充分激活。夹杂几何越复杂，声波传播路径越曲折，能量束缚和耗散效应也越明显。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-018.jpg",
                "caption": "图16 Design-1 的二维与三维模型吸声系数对比，以及四种设计在不同频率下的声压场与声速场"
            }
        ],
    },
    {
        "level": 1,
        "title": "5 实验验证",
        "paragraphs": [
            "为了验证逆向设计所得 MMPP 的可制造性与有效性，作者选取 Design-1 制作实物样件，并在测量频率范围为 60-1600 Hz 的方形阻抗管中，基于双传声器法测量其吸声系数。由于阻抗管边长为 120 mm，而原始设计单元的双拼尺寸为 100×100 mm^2，因此作者对单元进行了适当扩展，并将最终试样尺寸设为 119×119 mm^2，以便放置和取出。",
            "实物样件主要由顶部穿孔板、边缘刚性盖板、多孔基体以及内部刚性夹杂四部分组成。其中，穿孔板、盖板和刚性夹杂采用树脂并通过 SLA 3D 打印制造，打印精度为 0.01 mm；多孔材料则为三聚氰胺泡沫，并通过冲压形成所需形状。为确保实验可靠性，作者共重复测试三次。",
            "实验结果与仿真结果整体趋势一致，验证了数值分析的准确性。实验与仿真之间的偏差主要来自 3D 打印公差、装配过程中的配合误差，以及为支撑试样而额外引入盖板所带来的频率偏移。尽管如此，实验仍证明该逆向设计方案具备良好的可加工性和实际吸声效果。"
        ],
        "figures": [
            {
                "path": IMAGE_DIR / "img-019.jpg",
                "caption": "图17 双传声器法实验原理、实验环境、样件装配过程及实验与仿真结果对比"
            }
        ],
    },
    {
        "level": 1,
        "title": "6 结论",
        "paragraphs": [
            "本文提出了一种面向带穿孔板多孔超材料的 CGAN 逆向设计方法，旨在拓宽结构设计空间并实现高效宽频吸声。作者在先验知识指导下，通过控制函数与随机参数组合生成多类型 MMPP 图样，并建立了覆盖较宽吸声频段的数据集。训练后的 CGAN 可以在高自由度空间中进行面向目标吸声谱的快速搜索，无需重复进行大量正向仿真。",
            "测试结果表明，该方法在测试集上取得了较高的逆向设计精度与图样有效率，既能够生成接近训练样本分布的结构，也能够组合出超出原始数据集的新型方案。进一步分析还发现，低频吸声倾向于对应尺寸较小、浅层埋设的夹杂，而高频吸声则更有利于采用尺寸更大、埋设更深且填充率更高的夹杂。",
            "与仅关注吸收峰增强的已有研究不同，本文方法能够支持用户自定义吸声谱整形，并在 360-3000 Hz 的有效频段内生成多组候选方案。结合数值分析与实验验证，作者揭示了微穿孔板共振、夹杂局部共振、高阶模态以及声波干涉共同作用下的协同吸声机理。",
            "总体而言，该逆向设计方法突破了高维参数正向设计的计算瓶颈，在兼顾可制造性的前提下实现了多样化、创造性的结构设计；更重要的是，它还能在数据驱动框架中学习结构与性能之间的物理规律，因此既可作为工程设计工具，也可作为科研中探索结构-性能关系的辅助方法。"
        ],
    },
]


def set_run_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def style_paragraph(paragraph, font_name: str, size_pt: int, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, font_name, size_pt, bold)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, "宋体", 12, False)


def add_heading(doc: Document, title: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 16
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 15
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 14
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(title)
    set_run_font(run, "黑体", size, True)


def add_blank_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        return

    NORMALIZED_IMAGE_DIR.mkdir(parents=True, exist_ok=True)

    normalized_path = NORMALIZED_IMAGE_DIR / f"{image_path.stem}.png"
    with Image.open(image_path) as img:
        if img.mode not in ("RGB", "RGBA"):
            img = img.convert("RGB")
        width_px, height_px = img.size
        img.save(normalized_path, format="PNG")

    max_width_cm = 14.5
    width_cm = min(max_width_cm, width_px / 110)
    if height_px > width_px * 1.35:
        width_cm = min(11.5, width_cm)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run()
    run.add_picture(str(normalized_path), width=Cm(width_cm))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    cap.paragraph_format.line_spacing = Pt(18)
    run = cap.add_run(caption)
    set_run_font(run, "宋体", 10, False)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(0)
    title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    title_p.paragraph_format.line_spacing = Pt(20)
    title_run = title_p.add_run(TITLE)
    set_run_font(title_run, "黑体", 18, True)

    add_blank_line(doc)

    for block in CONTENT:
        add_heading(doc, block["title"], block["level"])
        for paragraph in block["paragraphs"]:
            add_body_paragraph(doc, paragraph)
        for figure in block.get("figures", []):
            add_figure(doc, figure["path"], figure["caption"])
        add_blank_line(doc)

    for _ in range(4):
        add_blank_line(doc)

    source_title = doc.add_paragraph()
    source_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    source_title.paragraph_format.line_spacing = Pt(20)
    source_title.paragraph_format.space_before = Pt(0)
    source_title.paragraph_format.space_after = Pt(0)
    title_run = source_title.add_run("译文原文出处：")
    set_run_font(title_run, "黑体", 12, True)

    source_p = doc.add_paragraph()
    source_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    source_p.paragraph_format.line_spacing = Pt(20)
    source_p.paragraph_format.space_before = Pt(0)
    source_p.paragraph_format.space_after = Pt(0)
    source_run = source_p.add_run(SOURCE)
    set_run_font(source_run, "宋体", 12, False)

    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
