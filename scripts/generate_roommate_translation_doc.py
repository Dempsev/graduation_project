from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_DIR = Path(r"D:\graduation_project\coad\output\doc")
OUTPUT_PATH = OUTPUT_DIR / "毕业设计文献翻译-反推装置.docx"
PAGE_DIR = Path(r"D:\graduation_project\coad\tmp\roommate_pages")
CROP_DIR = Path(r"D:\graduation_project\coad\tmp\roommate_crops")

TITLE = "飞机发动机两种反推装置的分析与仿真"
SOURCE = (
    "Tian F, He J, Xiong Y, Zhao Y. Analysis and simulation on two types of thrust reversers in an aircraft engine[C]. "
    "MATEC Web of Conferences, 2017, 119: 01012. DOI: 10.1051/matecconf/201711901012."
)


CONTENT = [
    {
        "level": 1,
        "title": "摘要",
        "paragraphs": [
            "随着新型复合材料和先进制造技术的发展，发动机短舱正在不断引入创新性的工程方案，例如一体化推进系统、通过整体成型工艺制造的碳纤维复合材料内壁等。这些技术不仅能够降低燃油消耗，还能够减小发动机噪声。在这一背景下，先进短舱中出现了一种 O 型外涵道反推装置验证结构，其复合材料外形呈“O”字形，而传统结构则通常为“D”型外涵道。",
            "本文围绕最新的 O-duct 与传统 D-duct 两种反推装置展开对比研究，重点通过数值建模与仿真方法分析二者之间的差异。为了突出反推装置工作过程中的定量特征，作者主要利用 CATIA 与数字样机（DMU）模块，对两种不同反推装置在收起状态和展开状态下的运动过程进行模拟。",
            "在完成结构重量比较之后，作者建立了挡流门相关机构的设计模型，并进一步进行了机构运动学分析与仿真。研究结果表明，结构简化以及多处接口的消除能够带来明显的减重效果；O 型外涵道有助于改善发动机内部气流流动；而 D 型外涵道则在成本效益和可维护性方面仍具有较好的优势。"
        ],
    },
    {
        "level": 1,
        "title": "1 引言",
        "paragraphs": [
            "现代大型客机在最大起飞重量和起飞速度方面不断提升，因此无论起飞还是着陆都对跑道长度提出了更高要求。面对日益增长的机场流量以及各种紧急工况，必须依靠可靠的减速系统来有效缩短着陆距离或中断起飞滑跑距离。当前民航飞机常见的减速方式包括机轮制动、扰流板空气制动、反推装置、阻力伞等。虽然现代大型客机大多配备了反推装置，但其设计会直接影响发动机短舱重量、机翼气动性能以及巡航效率，从而在一定程度上决定总体运行与维护成本。",
            "文中指出，GE 的研究认为反推装置重量约占整个短舱重量的 30%。对于风扇直径超过 2.5 米的大涵道比涡扇发动机而言，反推装置引起的泄漏和压降还会导致飞行比油耗增加 0.5% 到 1.0%。波音公司的经济性评估也显示，B767 使用反推装置每年需要额外增加 12.5 万美元左右的费用。",
            "尽管如此，反推装置依然不可或缺，因为它相较其他制动方式具有明显优势。反推在飞机减速直至停止过程中都能持续发挥作用，并且不像轮胎制动那样容易受到湿跑道、冰雪天气等外界条件影响，因此能够在恶劣环境着陆、迫降或中断起飞等情况下提供可靠的减速能力，从而提升飞行安全性。",
            "传统反推装置包括靶式、蛤壳式和冷气流式等类型。其中，级联式或冷气流式反推结构完整性较好、适用性较强，且能较稳定地提供相当于发动机最大推力 60% 到 70% 的反推效果，因此尤其适用于涡扇发动机。作者强调，挡流门机构是反推装置的重要组成部分，会直接影响反推力大小及系统效率，其运动方式、开启时间以及最终偏转角度都与机构设计密切相关。",
            "在结构形式方面，D-duct 系统通常由左右两个反推半体组成，并通过连接件和内侧 V 型构件固定在发动机吊架周围；其平移罩由内外套筒组成，外套筒构成外轮廓，内套筒形成风扇涵道外壁并承载挡流门。与之相比，O-duct 的显著特征是采用了接近 330 度的一体化碳纤维复合材料内壁，不再像 D-duct 那样分为两半，这使得平移罩结构更加整体化，也减少了锁扣等部件带来的重量。",
            "此外，O-duct 取消了暴露在外涵道中的阻力拉杆，改用布置在级联栅下方的更复杂挡流门机构。平移罩在电动作动筒驱动下向后移动的同时，也会带动挡流门转动到工作位置，从而改变气流方向并产生反推。基于这一差异，本文将重点对两种结构的模型建立、机构运动学与动力学特征进行分析。"
        ],
    },
    {
        "level": 1,
        "title": "2 方法与流程",
        "paragraphs": [
            "本文的建模和仿真工作以 CFM 系列发动机反推装置的结构模型为基础，在 CATIA 中建立两种反推装置的三维模型。为了保证仿真完整性与分析结果的可靠性，作者并非只单独建立反推器局部模型，而是将整个短舱系统一并建模，以提高反推器模型质量。",
            "作者指出，反推装置虽存在若干典型结构形式，但其工程研究价值主要体现在大型运输机应用场景中，而挡流门机构的布置形式又并不完全相同。因此，本文选择最常见的 D-duct 和较新的 O-duct 作为比较对象，重点聚焦于仿真模型建立和运动学分析。该研究方法也可以推广到其他类似反推系统的设计分析中。",
            "根据 CATIA 中建立的几何模型，作者在 DMU 中开展运动学仿真，并进一步在 ABAQUS 中进行动力学分析。运动学分析主要包括运动副规律设定、位移与速度参数测量以及相关特征分析；随后再将模型导入 ABAQUS，对挡流门机构在工作过程中的受载谱进行动态响应分析。"
        ],
    },
    {
        "level": 1,
        "title": "3 结构模型",
        "paragraphs": [
            "作者首先给出了短舱若干关键几何参数的估算公式，包括进气道唇口直径、主整流罩最大高度、主整流罩长度、风扇出口处主整流罩直径以及燃气发生器整流罩出口直径等。在此基础上得到对应的计算结果，并结合典型短舱外形示意图，为后续两类反推装置建立统一尺寸基准，以保证比较过程尽可能公平。",
            "为了使重量比较更有可比性，作者尽量保持两类反推装置的附属结构尺寸一致，并对模型进行适当简化。总体思路是：在不破坏主要受力特征和运动特性的前提下，通过合理简化降低建模与仿真复杂度。"
        ],
        "figures": [
            {"type": "crop", "page": "page-3.png", "box": (40, 520, 1060, 1400), "caption": "图1 典型短舱结构及主要几何参数计算结果"}
        ],
    },
    {
        "level": 2,
        "title": "3.1 D 型外涵道反推装置",
        "paragraphs": [
            "D-duct 反推装置的平移罩组件位于发动机后部，在飞行过程中保持前收状态；当反推被驱动时，平移罩向后移动，带动挡流门转动，封堵风扇涵道与反推平移罩之间的气流通道，并使级联栅暴露出来。随后气流经级联栅导向前方，形成反向推力。",
            "基于 D-duct 的结构特征，作者在 CATIA 中建立了主要部件模型，包括平移罩、内壁、三个液压作动器及两个同步轴、六个级联栅、五个挡流门、五个挡流门阻力拉杆、位于吊架上的反推开启作动器、扭矩盒、风扇涵道内壁、检修门以及上下滑轨等。根据模型测量结果，作者还给出了反推装置的主要外形参数和不同工作状态下平移罩的行程。"
        ],
        "figures": [
            {"type": "crop", "page": "page-4.png", "box": (60, 620, 1030, 1120), "caption": "图2 D 型外涵道反推装置收起状态与展开状态"}
        ],
    },
    {
        "level": 2,
        "title": "3.2 O 型外涵道反推装置",
        "paragraphs": [
            "O-duct 在总体构成上与 D-duct 类似，也由平移罩、内壁板、作动器、挡流门机构、级联栅、扭矩盒、检修门和滑轨等部分组成。但由于平移罩由左右两半改为一体式结构，其挡流门数量增加，对应级联栅数量也更多，同时不再需要张力锁扣。这一变化不仅使短舱整体结构更紧凑，也带来了更好的集成效果。",
            "更关键的是，O-duct 的挡流门机构取消了传统的阻力拉杆，从而有助于降低外涵道内的附加气动阻力。作者认为，从结构对比结果来看，O-duct 具有更完整、更连续的气流流通区域，意味着在外涵道中气流分叉更少、流动更顺畅。不过，新技术的引入也伴随着维护性和经济性方面的不确定因素，例如复合材料制造与复杂新机构在维修便利性上可能不如传统结构。"
        ],
        "figures": [
            {"type": "crop", "page": "page-5.png", "box": (30, 50, 1035, 510), "caption": "图3 D 型与 O 型外涵道反推装置横截面对比"}
        ],
    },
    {
        "level": 1,
        "title": "4 仿真与结果",
        "paragraphs": [
            "在完成结构建模之后，作者分别从数字样机仿真和结果分析两个方面比较两类反推装置。"
        ],
    },
    {
        "level": 2,
        "title": "4.1 DMU 仿真",
        "paragraphs": [
            "数字样机（DMU）是一种利用三维数字模型描述产品的工程方法，其主要目标是缩短研制周期、在设计阶段提前识别潜在问题、减少物理样机数量以及提高产品开发质量。本文利用 DMU 对反推装置机构进行运动学模拟，以研究其展开与收起过程中的运动关系。",
            "作者给出了 DMU 仿真的流程图，并说明仿真重点在于建立正确的运动副约束、设置合适的时间关系以及跟踪关键测量点的位移和速度变化规律。"
        ],
        "figures": [
            {"type": "crop", "page": "page-5.png", "box": (120, 540, 940, 1085), "caption": "图4 DMU 仿真流程图"}
        ],
    },
    {
        "level": 2,
        "title": "4.2 结果分析",
        "paragraphs": [
            "在结果分析中，作者对 O-duct 的挡流门机构进行了适度简化，建立运动副并设置与时间相关的驱动规律后，在 CATIA 中完成 DMU 仿真。仿真中，两种反推装置的展开时间都设为 1.5 s，并在工作位置附近保持约 10 s。通过仿真可以比较直观地观察挡流门的工作过程。",
            "以 O-duct 为例，每组挡流门机构除挡流门本体外主要由五个构件组成。为了获取运动参数曲线，作者在关键构件上设置测量点和参考点，CATIA 则输出对应的运动参数变化曲线。文中的速度曲线测点位于伸缩杆位置，用于分析机构展开过程中线速度的变化规律。"
        ],
        "figures": [
            {"type": "crop", "page": "page-6.png", "box": (40, 40, 1020, 530), "caption": "图5 O 型外涵道挡流门机构及其参考点线速度曲线"}
        ],
    },
    {
        "level": 1,
        "title": "5 动力学分析",
        "paragraphs": [
            "为了进一步评估两类反推装置在工作状态下的受力特征，作者又对反推力计算和挡流门机构动力学响应进行了分析。"
        ],
    },
    {
        "level": 2,
        "title": "5.1 反推力分析",
        "paragraphs": [
            "在获得级联栅出口流量和外涵道流量之后，作者依据反推气流出口角、喷流速度以及着陆速度等参数，建立了外涵道反推力的计算关系式。其中文中的出口角 α 是由级联栅几何形状决定的关键参数。",
            "根据工程设定，本文假定反推力约为发动机最大推力的 30%，作用在挡流门上的压强为 0.0035 MPa。这一结果随后被用作 ABAQUS 动力学仿真的载荷输入。"
        ],
    },
    {
        "level": 2,
        "title": "5.2 动力学仿真",
        "paragraphs": [
            "为了提高仿真效率并保证结果准确性，作者在保留必要结构特征的前提下，对挡流门机构模型进行了简化，并将 CATIA 数据导入 ABAQUS 进行网格划分与动力学求解。结构材料设定为铝合金，杨氏模量取 71 GPa，密度取 2800 kg/m^3。仿真假设两类挡流门具有相同的开启时间和反推偏转角，并忽略级联栅对受力的影响。",
            "位移云图表明，两种挡流门的最大位移都出现在挡流门边缘区域；应力云图则显示，最大应力集中在挡流门机构与平移罩连接的铰接位置。作者给出的计算结果显示，O-duct 挡流门机构的最大应力约为 91 MPa，几乎是 D-duct 对应值 57.1 MPa 的两倍。这说明 O-duct 在减重与气动改善方面具有潜力，但其机构受力水平也更高，需要在结构强度与可靠性设计上给予更多关注。"
        ],
        "figures": [
            {"type": "crop", "page": "page-6.png", "box": (40, 900, 1030, 1510), "caption": "图6 O 型与 D 型挡流门位移响应示意图"},
            {"type": "crop", "page": "page-7.png", "box": (35, 120, 1025, 640), "caption": "图7 O 型与 D 型挡流门应力云图"}
        ],
    },
    {
        "level": 1,
        "title": "6 结论",
        "paragraphs": [
            "本文通过在 CATIA 中建立两类反推装置模型，并结合 DMU 运动学仿真与 ABAQUS 动力学分析，对 O-duct 和 D-duct 进行了对比研究。研究表明，如果不考虑吊架因素，仅从反推器本体结构看，D-duct 因为包含较多张力锁扣和连接件而整体更重；O-duct 则由于结构简化和一体化设计具有更明显的减重潜力。",
            "从 DMU 仿真结果看，O-duct 要达到完全展开状态所需时间更长，而且其复杂挡流门机构在工作过程中可能引入更多运动协调问题；从动力学分析结果看，最大位移主要出现在挡流门边缘，而应力主要集中在平移罩与挡流门机构的连接位置。",
            "综合而言，O-duct 在减重和改善气流组织方面具有优势，但其复杂机构也带来了更高的应力和潜在设计风险；D-duct 虽然结构较重，却在成本和维护性方面更具现实工程优势。"
        ],
    },
]


def set_run_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_blank_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)


def add_heading(doc: Document, title: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 16
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 15
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 14
    run = p.add_run(title)
    set_run_font(run, "黑体", size, True)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, "宋体", 12, False)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "宋体", 10, False)


def crop_image(page_name: str, box: tuple[int, int, int, int]) -> Path:
    CROP_DIR.mkdir(parents=True, exist_ok=True)
    page_path = PAGE_DIR / page_name
    crop_name = f"{Path(page_name).stem}-{box[0]}-{box[1]}-{box[2]}-{box[3]}.png"
    crop_path = CROP_DIR / crop_name
    if not crop_path.exists():
        with Image.open(page_path) as img:
            cropped = img.crop(box)
            cropped.save(crop_path, format="PNG")
    return crop_path


def add_figure(doc: Document, figure: dict) -> None:
    image_path = crop_image(figure["page"], figure["box"])
    with Image.open(image_path) as img:
        w, h = img.size
    width_cm = min(14.5, w / 95)
    if h > w * 1.2:
        width_cm = min(10.5, width_cm)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, figure["caption"])


def build_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(TITLE)
    set_run_font(run, "黑体", 18, True)

    add_blank_line(doc)

    for block in CONTENT:
        add_heading(doc, block["title"], block["level"])
        for para in block["paragraphs"]:
            add_body_paragraph(doc, para)
        for figure in block.get("figures", []):
            add_figure(doc, figure)
        add_blank_line(doc)

    for _ in range(4):
        add_blank_line(doc)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run("译文原文出处：")
    set_run_font(run, "黑体", 12, True)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(SOURCE)
    set_run_font(run, "宋体", 12, False)

    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
