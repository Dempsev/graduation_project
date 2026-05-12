from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_DIR = Path(r"D:\graduation_project\coad\output\doc")
OUTPUT_PATH = OUTPUT_DIR / "毕业设计文献翻译-增材制造拓扑优化.docx"
PAGE_DIR = Path(r"D:\graduation_project\coad\tmp\third_pages")
CROP_DIR = Path(r"D:\graduation_project\coad\tmp\third_crops")

TITLE = "面向增材制造的设计：通过拓扑优化实现轻量化并支撑月球航天器开发"
SOURCE = (
    "Orme M E, Gschweitl M, Ferrari M, Madera I, Mouriaux F. Designing for Additive Manufacturing: "
    "Lightweighting Through Topology Optimization Enables Lunar Spacecraft[J]. Journal of Mechanical Design, "
    "2017, 139(10): 100905. DOI: 10.1115/1.4037304."
)

CONTENT = [
    {
        "level": 1,
        "title": "摘要",
        "paragraphs": [
            "本文提出了一套面向航天飞行合格件的端到端增材制造开发方法，并以一个由五个大型轻量化拓扑优化构件组成的发动机支撑系统为案例进行说明。该支撑系统将用于 SpaceIL 参与 Google Lunar XPrize 竞赛的月球着陆器，用来安装发动机。作者强调，为了让工业界能够更高效地采用拓扑优化与增材制造技术，开发流程中加入了前期设计探索步骤，以尽量减少后续数值计算工作量。",
            "该方法不仅包括拓扑优化本身，还考虑了许多拓扑优化算法中尚未直接体现的增材制造约束，例如构建方向选择、悬垂控制以及支撑结构最小化等问题。针对完成拓扑优化后的设计，研究团队采用粉末床激光熔化技术进行制造，并通过严格的测试、校核与验证流程完成构件开发。",
            "案例结果表明，增材制造与拓扑优化的结合能够显著减轻结构重量，并支持复杂航天结构的快速开发。与原始基线设计相比，最终五构件系统质量由 4.0 kg 降至 2.95 kg，体现了该方法在月球航天器轻量化方面的明显潜力。"
        ],
    },
    {
        "level": 1,
        "title": "1 引言",
        "paragraphs": [
            "通过数字化模型逐层制造金属结构件的工艺被称为增材制造（AM），目前已在航空航天、汽车、医疗和能源等领域受到越来越多关注。增材制造的价值在于它能够制造经拓扑优化后显著轻量化的结构件，能够将多个部件整合为一个整体件以提升体积利用率，能够制造高度复杂的几何结构，还可用于按需制造以应对零部件停产问题，并显著缩短从概念、设计、制造到验证交付之间的全流程时间。",
            "本文选取了一个具有代表性的案例：SpaceIL 参与 Google Lunar XPrize 竞赛时所使用的月球着陆器发动机安装结构。竞赛任务要求在 2017 年内实现无人航天器月面着陆，这对零部件的交付周期和轻量化水平都提出了很高要求。",
            "作者指出，增材制造特别适合这一挑战，原因主要有两点。其一，增材制造可以在部件概念提出后快速完成制造与合格件交付，有助于满足紧迫的发射时间节点；其二，它能够制造经拓扑优化获得的轻量化结构件，而这对航空航天系统而言尤其具有吸引力。"
        ],
    },
    {
        "level": 1,
        "title": "2 背景与整体流程",
        "paragraphs": [
            "增材制造通常是指利用数字信息逐层构建功能零件的技术体系。本文采用的是选择性激光熔化（SLM）工艺。作者说明，增材制造的一个重要意义在于它不需要模具或传统切削加工，因此非常适合制造拓扑优化后的复杂结构，例如中空结构、带内部冷却通道的结构、仿生有机形态结构以及填充晶格的轻量化结构等。",
            "为了稳定获得可重复的显微组织和力学性能，作者团队建立了一套覆盖设计、制造和测试的整体流程。该流程包括五个主要步骤：候选零件选择与概念开发、拓扑优化及面向增材制造的设计解释、有限元验证、带见证试样的增材制造，以及材料验证与力学性能鉴定。",
            "流程中还包含两个关键反馈环：一是拓扑优化与有限元验证之间的反馈，二是制造与材料/结构测试之间的反馈。前者的作用在于对拓扑优化结果进行人工后处理，以满足悬垂角、支撑最小化、装配细节和大尺寸零件分体制造等增材制造准则；后者则用于在材料和力学性能未达到要求时，通过调整工艺参数重新制造试样和构件，直到满足最终设计要求。"
        ],
        "figures": [
            {"page": "page-2.png", "box": (110, 40, 950, 250), "caption": "图1 高质量高可靠金属增材制造的整体流程图"}
        ],
    },
    {
        "level": 1,
        "title": "3 候选部件选择与拓扑优化设计",
        "paragraphs": [
            "本文选择的增材制造对象由四个支腿和一个中央连接毂组成，其功能是支撑航天器发动机并与隔热罩及航天器环形结构相连接。所有零件均采用 AlSi10Mg 铝合金制造。该组件需要承受较高载荷和较宽温度范围，同时还应具备低质量和高刚度，因此非常适合作为拓扑优化的候选对象。",
            "作者使用 Altair HyperWorks 14.0 进行拓扑优化，利用惩罚型等效各向同性材料方法，在满足给定载荷和几何要求的前提下确定材料最佳分布。设计空间由接口位置、保留体积以及非设计区域共同确定，其中保留体积用于避免与相邻航天器部件干涉，非设计区域则对应连接点等必须保留材料的区域。",
            "优化目标设定为在满足首阶固有频率大于 60 Hz、最大应力小于 115 MPa 的前提下尽可能减小构件质量。为了节省计算资源，拓扑优化阶段采用了较粗的网格，而后续有限元验证使用更细网格。初步设计探索表明，原本连续的支腿结构倾向于分裂成三条仿生分支，并且这些分支会在两个高度位置与中央连接毂汇接，这些趋势为后续缩小设计空间和进一步优化提供了依据。"
        ],
        "figures": [
            {"page": "page-2.png", "box": (55, 450, 1035, 1450), "caption": "图2 发动机支撑结构基线设计与初步拓扑优化探索结果"},
            {"page": "page-3.png", "box": (30, 510, 1035, 1470), "caption": "图3 设计趋势识别、缩减设计空间及最终分体设计空间示意图"}
        ],
    },
    {
        "level": 2,
        "title": "3.1 面向增材制造的设计解释",
        "paragraphs": [
            "在拓扑优化得到总体设计概念后，作者需要进一步判断大型结构应在何处、以何种方式分割成多个零件，以适配 EOS M290 设备的成形空间。最终设计被拆分为四个相同支腿和一个中央毂，并通过高配合精度的剪切螺栓连接。",
            "与此同时，设计还必须尽量实现自支撑结构，或在无法完全自支撑时将支撑数量降到最低。作者采用的一条经验规则是：相对于成形基板悬垂角大于 45° 的几何特征通常可以无需支撑直接打印。因此，零件构型不仅受到力学需求驱动，也受到打印方向和成形可制造性的显著影响。最终系统总质量降至 2.95 kg，相比原始 4.0 kg 基线设计实现了明显减重。"
        ],
        "figures": [
            {"page": "page-4.png", "box": (45, 35, 1025, 1475), "caption": "图4 四支腿与中央连接毂的最终总装结构、分体零件及连接建模方法"}
        ],
    },
    {
        "level": 1,
        "title": "4 有限元验证",
        "paragraphs": [
            "由于拓扑优化后的最终几何并不完全等同于算法直接输出结果，因此必须通过独立的有限元分析再次验证设计。本文所有分析均在 HyperWorks 14.0 平台完成，采用 Hypermesh 建模、OptiStruct 求解和 Hyperview 后处理。为了较准确反映应力集中，作者使用了相对较小的网格尺寸，整个模型约包含 100 万个单元和 30 万个节点。",
            "增材制造构件采用实体四面体单元建模，膨胀锥和隔热罩采用壳单元建模，发动机则采用足够详细的有限元模型以匹配供应商给出的模态结果。支腿在与航天器连接的四个脚位处施加全约束，以模拟与航天器结构的真实连接状态。零件之间的连接采用 RBE2、BUSH 和 BAR 元组合建模，以较保守的方式表示螺栓连接刚度，而螺栓预紧和接触压力则未显式纳入模型。",
            "考虑到增材制造仍属于相对较新的工艺，缺少成熟的航天历史数据库支撑，因此作者在常规航天安全系数之外额外引入了 1.5 的增材制造保守系数。结构同时承受准静态载荷、正弦振动和随机振动载荷。分析结果表明，在 xy 平面与 z 方向的等效静载条件下，结构最大应力均满足设计许用值要求，安全裕度分别为 0.04 和 0.21，因此按流程判断，该设计已经具备制造条件。"
        ],
        "figures": [
            {"page": "page-5.png", "box": (140, 250, 920, 1280), "caption": "图5 发动机支撑结构在 xy 激励和 z 激励下的有限元应力结果"}
        ],
    },
    {
        "level": 1,
        "title": "5 增材制造与试验验证",
        "paragraphs": [
            "本文所示总装由三块构建板打印完成：左侧构建板包含两个支腿，中间构建板包含连接毂，右侧构建板包含另外两个支腿。每块构建板上同时打印了拉伸试样和密度试样等见证件，其中与中央毂同板打印的竖直试样还被布置在本来需要支撑结构的区域，因此兼具工艺验证和支撑替代两种作用。",
            "材料验证阶段对试样进行了极限强度、屈服强度、延伸率和弹性模量测试，并对密度试样进行金相截面观察。试验结果显示，所有测得性能均超过预先设定的设计许用值，因此后续继续开展了对实际打印件的无损检测。",
            "在构件层面，作者使用 CT 扫描检测内部缺陷和孔隙，同时检查打印件与原始 CAD 模型之间的几何偏差。结果表明，未发现大于 400 微米的缺陷，五个构件的最大尺寸偏差约为 500 微米，均在可接受范围之内，且测得致密度达到 99.7%。动态结构试验则仍在继续进行，包括低级别正弦扫频、高级别正弦、随机和冲击试验；若结果不满足要求，则需返回制造步骤，重新调整工艺参数并再次打印和测试。"
        ],
        "figures": [
            {"page": "page-6.png", "box": (35, 30, 1035, 980), "caption": "图6 三块构建板上的发动机支撑组件、见证试样及装配后的发动机模型照片"}
        ],
    },
    {
        "level": 1,
        "title": "6 总结",
        "paragraphs": [
            "本文将一套从概念设计到零件合格认证的整体流程应用于一个计划在 2017 年执行月球飞行任务的五构件发动机支撑系统。研究重点不仅在于拓扑优化本身，更在于如何将拓扑优化结果转化为真正可增材制造、可装配、可验证的工程零件。",
            "最终获得的系统由四个仿生化支腿和一个中央毂组成，总质量从 4.0 kg 降至 2.95 kg，显著轻于原始基线设计。见证试样性能测试结果满足要求，CT 检测显示几何偏差和孔隙率都处于可接受范围内，结构试验也在进一步推进。",
            "作者据此认为，增材制造之所以能够成为面向航天飞行构件快速设计与制造的重要支撑技术，关键在于它一方面显著缩短了合格件交付周期，另一方面又使拓扑优化所得轻量化复杂结构真正具备制造可行性。配合整体流程控制，增材制造能够较可靠地支撑航天高质量结构件开发。"
        ],
    },
]


def set_run_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_blank_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)


def add_heading(doc: Document, title: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 16
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 15
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 14
    run = p.add_run(title)
    set_run_font(run, "黑体", size, True)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, "宋体", 12, False)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "宋体", 10, False)


def crop_image(page_name: str, box: tuple[int, int, int, int]) -> Path:
    CROP_DIR.mkdir(parents=True, exist_ok=True)
    page_path = PAGE_DIR / page_name
    crop_name = f"{Path(page_name).stem}-{box[0]}-{box[1]}-{box[2]}-{box[3]}.png"
    crop_path = CROP_DIR / crop_name
    if not crop_path.exists():
        with Image.open(page_path) as img:
            img.crop(box).save(crop_path, format="PNG")
    return crop_path


def add_figure(doc: Document, figure: dict) -> None:
    image_path = crop_image(figure["page"], figure["box"])
    with Image.open(image_path) as img:
        w, h = img.size
    width_cm = min(14.5, w / 95)
    if h > w * 1.15:
        width_cm = min(10.8, width_cm)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, figure["caption"])


def build_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(TITLE)
    set_run_font(run, "黑体", 18, True)

    add_blank_line(doc)

    for block in CONTENT:
        add_heading(doc, block["title"], block["level"])
        for para in block["paragraphs"]:
            add_body_paragraph(doc, para)
        for figure in block.get("figures", []):
            add_figure(doc, figure)
        add_blank_line(doc)

    for _ in range(4):
        add_blank_line(doc)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run("译文原文出处：")
    set_run_font(run, "黑体", 12, True)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(SOURCE)
    set_run_font(run, "宋体", 12, False)

    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
